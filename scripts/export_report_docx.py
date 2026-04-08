"""
Convert docs/TECHNICAL_REPORT.md to docs/TECHNICAL_REPORT.docx.

A minimal, self-contained Markdown → DOCX converter that handles
the subset of Markdown the report actually uses:
- ATX headings (#, ##, ###, ####)
- Paragraphs with inline emphasis (**bold**, *italic*, `code`)
- Bullet lists ("- ")
- Fenced code blocks (```)
- GitHub-flavored tables (pipe syntax with header separator)
- Horizontal rules (---)
- Blockquotes (">")

We do NOT shell out to pandoc because the submission machine doesn't
have pandoc installed. python-docx is a small pure-Python dependency
that's already in the project's dev group.

Usage:
    uv run python -m scripts.export_report_docx
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor

REPO_ROOT = Path(__file__).resolve().parent.parent
SRC = REPO_ROOT / "docs" / "TECHNICAL_REPORT.md"
DST = REPO_ROOT / "docs" / "TECHNICAL_REPORT.docx"


# ── Inline emphasis ───────────────────────────────────────────────

INLINE_RE = re.compile(
    r"(\*\*(?P<bold>[^*]+)\*\*)"
    r"|(\*(?P<ital>[^*]+)\*)"
    r"|(`(?P<code>[^`]+)`)"
)


def add_inline_runs(paragraph, text: str) -> None:
    """Append text to `paragraph`, honoring bold/italic/code spans."""
    i = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > i:
            paragraph.add_run(text[i : match.start()])
        if match.group("bold"):
            run = paragraph.add_run(match.group("bold"))
            run.bold = True
        elif match.group("ital"):
            run = paragraph.add_run(match.group("ital"))
            run.italic = True
        elif match.group("code"):
            run = paragraph.add_run(match.group("code"))
            run.font.name = "Menlo"
            run.font.size = Pt(9)
        i = match.end()
    if i < len(text):
        paragraph.add_run(text[i:])


# ── Block dispatcher ──────────────────────────────────────────────

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")
TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$")


def split_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def convert(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()

    # Base font and page margins
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    i = 0
    in_code = False
    code_buffer: list[str] = []

    while i < len(lines):
        line = lines[i]

        # Fenced code block start/end
        if line.strip().startswith("```"):
            if in_code:
                # flush
                code_para = doc.add_paragraph()
                code_run = code_para.add_run("\n".join(code_buffer))
                code_run.font.name = "Menlo"
                code_run.font.size = Pt(9)
                code_para.paragraph_format.left_indent = Inches(0.25)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == "---":
            doc.add_paragraph().add_run("─" * 60)
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Heading
        h = HEADING_RE.match(line)
        if h:
            level = len(h.group(1))
            content = h.group(2).strip()
            para = doc.add_heading(level=min(level, 4))
            add_inline_runs(para, content)
            i += 1
            continue

        # Blockquote (report only has one)
        if line.lstrip().startswith(">"):
            content = line.lstrip()[1:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_inline_runs(p, content)
            i += 1
            continue

        # Table detection: two consecutive non-blank lines where
        # the second matches the separator pattern.
        if (
            "|" in line
            and i + 1 < len(lines)
            and TABLE_SEP_RE.match(lines[i + 1])
        ):
            header = split_table_row(line)
            i += 2  # skip header + separator
            rows: list[list[str]] = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                rows.append(split_table_row(lines[i]))
                i += 1
            tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
            tbl.style = "Light Grid Accent 1"
            for col_idx, cell_text in enumerate(header):
                cell = tbl.rows[0].cells[col_idx]
                cell.text = ""
                add_inline_runs(cell.paragraphs[0], cell_text)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for row_idx, row in enumerate(rows, start=1):
                for col_idx in range(len(header)):
                    cell = tbl.rows[row_idx].cells[col_idx]
                    cell.text = ""
                    value = row[col_idx] if col_idx < len(row) else ""
                    add_inline_runs(cell.paragraphs[0], value)
            continue

        # Bullet list
        b = BULLET_RE.match(line)
        if b:
            para = doc.add_paragraph(style="List Bullet")
            add_inline_runs(para, b.group(1))
            i += 1
            continue

        # Numbered list ("1. ", "2. ", ...)
        m = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if m:
            para = doc.add_paragraph(style="List Number")
            add_inline_runs(para, m.group(1))
            i += 1
            continue

        # Plain paragraph — accumulate until blank line
        para_lines = [line]
        j = i + 1
        while (
            j < len(lines)
            and lines[j].strip()
            and not HEADING_RE.match(lines[j])
            and not lines[j].lstrip().startswith(("-", "*", ">", "|", "```"))
            and not re.match(r"^\s*\d+\.\s+", lines[j])
        ):
            para_lines.append(lines[j])
            j += 1
        para = doc.add_paragraph()
        add_inline_runs(para, " ".join(l.strip() for l in para_lines))
        i = j

    doc.save(out_path)
    print(f"Wrote {out_path} ({out_path.stat().st_size:,} bytes)")


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"Missing source: {SRC}")
    convert(SRC, DST)


if __name__ == "__main__":
    main()
